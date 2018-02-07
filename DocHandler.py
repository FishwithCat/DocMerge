# coding=utf-8
import os
from docx import Document
from win32com import client as wc


class DocHandler(object):

    @staticmethod
    def save_as_docx(originPath, targetPath):
        print 'start transcoding...'
        word = wc.Dispatch('Word.Application')
        originPath = os.path.abspath(originPath)
        targetPath = os.path.abspath(targetPath)
        if not os.path.exists(targetPath):
            os.mkdir(targetPath)
        for fileName in os.listdir(originPath):
            filepath = os.path.join(originPath, fileName)
            doc = word.Documents.Open(unicode(filepath, "gb2312"))
            fileName = fileName.replace('.doc', '.docx')
            doc.SaveAs(os.path.join(targetPath, fileName), 12, False, "", True, "", False, False, False, False)
            doc.Close()
        word.Quit()

    @staticmethod
    def merge_file(originPath, mergeFilePath):
        print 'start merge...'
        originPath = os.path.abspath(originPath)
        files = os.listdir(originPath)
        document = Document()
        for fileName in files:
            print fileName.decode('gb2312')
            filepath = os.path.join(originPath, fileName)
            originFile = Document(filepath)
            for paragraph in originFile.paragraphs:
                document.add_paragraph(paragraph.text)
        document.save(mergeFilePath)

    @staticmethod
    def clear_path(path):
        if not os.path.exists(path):
            return
        for fileName in os.listdir(os.path.abspath(path)):
            os.remove(os.path.join(path, fileName))
